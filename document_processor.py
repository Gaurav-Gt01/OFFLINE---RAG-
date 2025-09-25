from typing import List
import fitz  # PyMuPDF
import docx
from PIL import Image
import pytesseract

def process_pdf(file) -> str:
    """
    Extract text from a PDF file using PyMuPDF.
    Accepts a file-like object or a filepath.
    """
    # If file is a BytesIO or similar
    if hasattr(file, 'read'):
        data = file.read()
        doc = fitz.open(stream=data, filetype="pdf")
    else:
        doc = fitz.open(file)
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def process_docx(file) -> str:
    """
    Extract text from a DOCX file using python-docx.
    Accepts a file-like object or a filepath.
    """
    if hasattr(file, 'read'):
        doc = docx.Document(file)
    else:
        doc = docx.Document(file)
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    return "\n".join(text)

def process_image(file) -> str:
    """
    Extract text from an image file using pytesseract OCR.
    Accepts a file-like object or a filepath.
    """
    if hasattr(file, 'read'):
        image = Image.open(file)
    else:
        image = Image.open(file)
    text = pytesseract.image_to_string(image)
    return text

def split_text(text: str, chunk_size: int = 1000, overlap: int = 100) -> List[str]:
    """
    Split text into chunks of specified size with overlap.
    """
    if not text:
        return []
    chunks = []
    start = 0
    length = len(text)
    while start < length:
        end = start + chunk_size
        chunk = text[start:end]
        chunks.append(chunk)
        start += chunk_size - overlap
    return chunks
